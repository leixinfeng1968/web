# 使用python-docx库生成Word文档
# 首先安装库：pip install python-docx

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建文档对象
doc = Document()

# 设置文档字体
def set_font(run, font_name='微软雅黑', size=12):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

# 添加标题
heading = doc.add_heading('用HTML制作动态交互初中数学课件指南', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(heading.runs[0], '微软雅黑', 24)

# 添加副标题
doc.add_paragraph()
subtitle = doc.add_paragraph('初中数学动态课件制作规范与流程')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(subtitle.runs[0], '微软雅黑', 16)

# 添加作者信息
doc.add_paragraph()
author_info = doc.add_paragraph('文档版本：V1.0  |  最后更新：2023年11月')
author_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(author_info.runs[0], '宋体', 10)

# 第一部分：用豆包制作HTML交互动态演示课件的规范
doc.add_page_break()
section1 = doc.add_heading('第一部分：用豆包制作HTML交互动态演示课件的规范', level=1)
set_font(section1.runs[0], '微软雅黑', 18)

# 1.1 提示词规范
doc.add_heading('1.1 提示词规范', level=2)
set_font(doc.paragraphs[-1].runs[0], '微软雅黑', 14)

# 添加提示词规范内容
p1 = doc.add_paragraph()
p1.add_run('豆包制作课件的标准提示词格式为：').bold = True
set_font(p1.runs[0], '微软雅黑', 12)

# 添加提示词示例
p2 = doc.add_paragraph()
p2.add_run('我是初中数学老师，请给我用HTML设计一个').bold = True
p2.add_run('【课件名称】').italic = True
p2.add_run('的动态交互课件。')
for run in p2.runs:
    set_font(run, '微软雅黑', 12)

# 添加提示词说明
doc.add_paragraph('说明：')
doc.add_paragraph('• 提示词应包含课件的核心功能和交互需求')
doc.add_paragraph('• 明确说明目标学生群体（七年级/八年级/九年级）')
doc.add_paragraph('• 指出需要展示的数学概念和知识点')
doc.add_paragraph('• 描述期望的交互方式和动画效果')

# 为以上段落设置字体
for i in range(-5, 0):
    for run in doc.paragraphs[i].runs:
        set_font(run, '微软雅黑', 12)

# 1.2 代码提交规范
doc.add_heading('1.2 代码提交规范', level=2)
set_font(doc.paragraphs[-1].runs[0], '微软雅黑', 14)

# 添加代码提交规范步骤
steps = [
    '豆包生成代码后，先在豆包内置的预览功能中查看效果',
    '如果发现不合适的地方，继续向豆包提出修改建议，直到满意为止',
    '复制豆包生成的完整满意代码',
    '打开微信小程序中的「腾讯文档」',
    '新建一个腾讯文档',
    '将复制的代码粘贴到腾讯文档中',
    '将腾讯文档导出为Word格式',
    '将Word文档发送到电脑桌面',
    '在电脑桌面上新建一个文本文档（.txt文件）',
    '打开新建的文本文档',
    '将Word文档中的代码复制并粘贴到文本文档中',
    '保存文本文档',
    '将文本文档的后缀名从.txt改为.html',
    '双击修改后的HTML文件即可打开并使用课件'
]

for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    for run in p.runs:
        set_font(run, '微软雅黑', 12)

# 1.3 代码格式要求
doc.add_heading('1.3 代码格式要求', level=2)
set_font(doc.paragraphs[-1].runs[0], '微软雅黑', 14)

# 添加代码格式要求内容
format_requirements = [
    '代码必须包含完整的HTML结构（DOCTYPE、html、head、body标签）',
    '使用UTF-8字符编码（在head标签中添加<meta charset="UTF-8">）',
    '代码中应包含适当的注释，解释关键功能和逻辑',
    '确保代码在主流浏览器（Chrome、Firefox、Edge）中能正常运行',
    '界面设计应简洁明了，突出数学内容',
    '交互按钮和控件应具有清晰的文字说明',
    '动画效果应流畅，不影响数学内容的展示'
]

for req in format_requirements:
    p = doc.add_paragraph(f'• {req}')
    for run in p.runs:
        set_font(run, '微软雅黑', 12)

# 第二部分：课件制作示例
doc.add_heading('第二部分：课件制作示例', level=1)
set_font(doc.paragraphs[-1].runs[0], '微软雅黑', 18)

# 添加示例说明
p = doc.add_paragraph('以下以「九年级下册第三章圆心角性质定理」课件为例，展示完整的课件制作流程：')
set_font(p.runs[0], '微软雅黑', 12)

# 2.1 示例提示词
doc.add_heading('2.1 示例提示词', level=2)
set_font(doc.paragraphs[-1].runs[0], '微软雅黑', 14)

p = doc.add_paragraph()
p.add_run('提示词：').bold = True
set_font(p.runs[0], '微软雅黑', 12)

p = doc.add_paragraph('我是初中数学老师，请给我用HTML设计一个「九年级下册第三章圆心角性质定理」的动态交互课件，主要功能包括：')
for run in p.runs:
    set_font(run, '微软雅黑', 12)

# 添加提示词的具体功能要求
features = [
    '1. 展示两个圆心角（∠AOB和∠COD），分别用不同颜色区分',
    '2. 提供三个滑动条控制：蓝色角大小、绿色角大小、绿色角位置',
    '3. 动态显示弦长、角度等信息',
    '4. 当两个圆心角相等时，移动绿色角可以演示重合效果',
    '5. 展示圆心角性质定理及其推论',
    '6. 添加交互提示，指导学生操作',
    '7. 界面设计美观，符合初中数学教学需求'
]

for feature in features:
    p = doc.add_paragraph(feature)
    for run in p.runs:
        set_font(run, '微软雅黑', 12)

# 2.2 课件功能实现
doc.add_heading('2.2 课件功能实现', level=2)
set_font(doc.paragraphs[-1].runs[0], '微软雅黑', 14)

# 添加功能实现说明
p = doc.add_paragraph('基于上述提示词，豆包生成了完整的HTML代码，实现了以下核心功能：')
set_font(p.runs[0], '微软雅黑', 12)

# 添加功能列表
implementations = [
    '• 动态图形展示：包含圆、圆心、点、半径、弦、弧等几何元素',
    '• 交互控制：三个滑动条和输入框，可精确控制角度和位置',
    '• 实时反馈：显示角度值、弦长等数据',
    '• 定理展示：显示圆心角性质定理及其推论',
    '• 验证机制：当两角相等且位置相同时，显示重合提示',
    '• 响应式设计：适配不同屏幕尺寸'
]

for impl in implementations:
    p = doc.add_paragraph(impl)
    for run in p.runs:
        set_font(run, '微软雅黑', 12)

# 2.3 操作说明
doc.add_heading('2.3 操作说明', level=2)
set_font(doc.paragraphs[-1].runs[0], '微软雅黑', 14)

# 添加操作步骤
operations = [
    '1. 调整前两个滑动条，设置蓝色角和绿色角的大小',
    '2. 当两个角的角度相等时，界面会显示提示信息',
    '3. 拖动第三个滑动条，移动绿色角的位置',
    '4. 观察绿色角与蓝色角的重合情况',
    '5. 阅读显示的圆心角性质定理及其推论'
]

for i, op in enumerate(operations, 1):
    p = doc.add_paragraph(op)
    for run in p.runs:
        set_font(run, '微软雅黑', 12)

# 2.4 教学应用
doc.add_heading('2.4 教学应用', level=2)
set_font(doc.paragraphs[-1].runs[0], '微软雅黑', 14)

p = doc.add_paragraph('该课件可用于九年级下册第三章《圆心角性质定理》的教学，帮助学生：')
set_font(p.runs[0], '微软雅黑', 12)

applications = [
    '• 直观理解圆心角与弦、弧的关系',
    '• 通过交互操作验证圆心角性质定理',
    '• 加深对定理推论的理解',
    '• 培养几何直观和空间想象力'
]

for app in applications:
    p = doc.add_paragraph(app)
    for run in p.runs:
        set_font(run, '微软雅黑', 12)

# 保存文档
doc.save('用HTML制作动态交互初中数学课件指南.docx')

print('文档已生成：用HTML制作动态交互初中数学课件指南.docx')
